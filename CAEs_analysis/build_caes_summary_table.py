
"""
Creates a Word summary table of CAE matching results across thresholds, using CSV summary files.
"""
import pandas as pd
from docx import Document
from pathlib import Path

settings = [
    ("0.3", "0_3"),
    ("0.4", "0_4"),
    ("0.5", "0_5"),
]

dfs = {}
for disp, suffix in settings:
    csv_path = Path(f"cae_summary_{suffix}.csv")
    if not csv_path.is_file():
        raise FileNotFoundError(f"Missing {csv_path}")
    df = pd.read_csv(csv_path).set_index("target")[["matched", "distorted", "no_match"]]
    dfs[disp] = df

target_order = [
    '1_ABAHIW',  '2_ABAKIZ',   '3_ABADOX',  '4_ABABIP',   '5_GASQOK',  '6_ABEKIE',
    '7_NIWPUE01','8_ABEKIF',   '9_APUFEX',  '10_ABEHAU', '11_TITTUO','12_EGEYOG',
    '13_ABOBUP','14_XIDTOW',  '15_ACNCOB10','16_TACXUQ','17_ACAZFE','18_NIVHEJ',
    '19_ADUPAS','20_DAJLAC',  '21_OFOWIS','22_CATSUL','23_HESMUQ01','24_GUDQOL',
    '25_ABEVAG','26_AKOQOH', '27_ADARUT','28_AFECIA','29_ACOVUL','30_AFIXEV',
    '31_ABAYAF','32_RULJAM'
]

df03 = dfs["0.3"]
totals = df03["matched"].add(df03["distorted"], 0).add(df03["no_match"], 0).astype(int)

doc = Document()
doc.add_heading("CAE Matching Summary Across Thresholds", level=1)

n_cols = 2 + 3 * len(settings)
table = doc.add_table(rows=1 + len(target_order), cols=n_cols)
table.style = "Light List Accent 1"

hdr = table.rows[0].cells
hdr[0].text = "Target"
hdr[1].text = "Total CAEs"
for i, (disp, _) in enumerate(settings):
    base = 2 + 3*i
    hdr[base  ].text = f"{disp} matched"
    hdr[base+1].text = f"{disp} distorted"
    hdr[base+2].text = f"{disp} no_match"

for ridx, numbered in enumerate(target_order, start=1):
    _, bare = numbered.split("_", 1)
    row = table.rows[ridx].cells

    row[0].text = numbered
    row[1].text = str(totals.get(bare, ""))

    for i, (disp, _) in enumerate(settings):
        df = dfs[disp]
        base = 2 + 3*i
        if bare in df.index:
            m, d, nm = df.loc[bare, ["matched", "distorted", "no_match"]]
            row[base  ].text = str(int(m))
            row[base+1].text = str(int(d))
            row[base+2].text = str(int(nm))
        else:
            row[base  ].text = ""
            row[base+1].text = ""
            row[base+2].text = ""

out = Path("cae_summary_table.docx")
doc.save(out)
print(f"✅  Written → {out.resolve()}")
