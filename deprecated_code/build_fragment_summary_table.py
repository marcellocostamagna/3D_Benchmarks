#!/usr/bin/env python3
import pandas as pd
from docx import Document
from pathlib import Path

# 1) Define thresholds and file‐suffixes
settings = [
    ("0.3", "0_3"),
    ("0.4", "0_4"),
    ("0.5", "0_5"),
]

# 2) Load each CSV into a dict of DataFrames, indexed by bare target ID
dfs = {}
for disp, suffix in settings:
    csv_path = Path(f"fragment_summary_{suffix}.csv")
    if not csv_path.is_file():
        raise FileNotFoundError(f"Missing {csv_path}")
    df = pd.read_csv(csv_path).set_index("target")[["matched", "fallback", "no_match"]]
    dfs[disp] = df

# 3) Your exact numeric‐prefix target order
target_order = [
    '1_ABAHIW',  '2_ABAKIZ',   '3_ABADOX',  '4_ABABIP',   '5_GASQOK',  '6_ABEKIE',
    '7_NIWPUE01','8_ABEKIF',   '9_APUFEX',  '10_ABEHAU', '11_TITTUO','12_EGEYOG',
    '13_ABOBUP','14_XIDTOW',  '15_ACNCOB10','16_TACXUQ','17_ACAZFE','18_NIVHEJ',
    '19_ADUPAS','20_DAJLAC',  '21_OFOWIS','22_CATSUL','23_HESMUQ01','24_GUDQOL',
    '25_ABEVAG','26_AKOQOH', '27_ADARUT','28_AFECIA','29_ACOVUL','30_AFIXEV',
    '31_ABAYAF','32_RULJAM'
]

# 4) Compute “Total fragments” once from the 0.3 DataFrame
df03 = dfs["0.3"]
# df03.index has bare IDs like "ABAHIW"
totals = df03["matched"].add(df03["fallback"], 0).add(df03["no_match"], 0).astype(int)

# 5) Create a Word document
doc = Document()
doc.add_heading("Fragment Matching Summary Across Thresholds", level=1)

# 6) Build the table
n_cols = 2 + 3 * len(settings)
table = doc.add_table(rows=1 + len(target_order), cols=n_cols)
table.style = "Light List Accent 1"

# 6a) Header
hdr = table.rows[0].cells
hdr[0].text = "Target"
hdr[1].text = "Total fragments"
for i, (disp, _) in enumerate(settings):
    base = 2 + 3*i
    hdr[base  ].text = f"{disp} matched"
    hdr[base+1].text = f"{disp} fallback"
    hdr[base+2].text = f"{disp} no_match"

# 6b) Fill rows
for ridx, numbered in enumerate(target_order, start=1):
    # split off the numeric prefix
    _, bare = numbered.split("_", 1)
    row = table.rows[ridx].cells

    # column 0 = "1_ABAHIW", etc.
    row[0].text = numbered
    # column 1 = total from 0.3
    row[1].text = str(totals.get(bare, ""))

    # fill each threshold
    for i, (disp, _) in enumerate(settings):
        df = dfs[disp]
        base = 2 + 3*i
        if bare in df.index:
            m, f, nm = df.loc[bare, ["matched", "fallback", "no_match"]]
            row[base  ].text = str(int(m))
            row[base+1].text = str(int(f))
            row[base+2].text = str(int(nm))
        else:
            row[base  ].text = ""
            row[base+1].text = ""
            row[base+2].text = ""

# 7) Save
out = Path("fragment_summary_table.docx")
doc.save(out)
print(f"✅  Written → {out.resolve()}")
