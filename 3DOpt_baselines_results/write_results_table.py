#!/usr/bin/env python3

import sys
from pathlib import Path
import pandas as pd
from docx import Document
import re

def find_results_csvs(folder):
    folder = Path(folder)
    for d1 in folder.iterdir():
        if d1.is_dir():
            for d2 in d1.iterdir():
                if d2.is_dir():
                    f = d2 / "results.csv"
                    if f.exists():
                        yield f

def clean_text(s):
    # Remove all leading/trailing whitespace (including Unicode, tabs, etc)
    return re.sub(r'^\s+|\s+$', '', str(s))

def write_results_table(results_folder):
    results_folder = Path(results_folder)
    csv_files = sorted(find_results_csvs(results_folder))
    if not csv_files:
        print(f"❌ No results.csv files found in {results_folder}")
        sys.exit(1)

    all_scores = {}
    n_runs = len(csv_files)
    for i, f in enumerate(csv_files):
        df = pd.read_csv(f)
        for _, row in df.iterrows():
            task = clean_text(row["task"])
            score = float(row["3DOpt_Score"])
            if task not in all_scores:
                all_scores[task] = [None] * n_runs
            all_scores[task][i] = score

    tasks_sorted = sorted(all_scores.keys(), key=lambda s: (int(s.split("_")[0]), s))

    doc = Document()
    doc.add_heading(f"3DOpt Baseline Results: {results_folder.name}", level=1)

    n_rows = 2 + len(tasks_sorted)
    n_cols = 1 + n_runs
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Light List Accent 1"

    # Header
    table.cell(0,0).text = "Task"
    run_cell = table.cell(0,1)
    run_cell.text = "Run"
    for col in range(2, n_cols):
        table.cell(0,col).text = ""
        run_cell = run_cell.merge(table.cell(0,col))
    table.cell(1,0).text = ""
    for j in range(n_runs):
        table.cell(1,j+1).text = clean_text(str(j+1))

    # Data
    for i, task in enumerate(tasks_sorted):
        row = table.rows[i+2].cells
        row[0].text = clean_text(task)
        for j, score in enumerate(all_scores[task]):
            cell_text = f"{score:.2f}" if score is not None else ""
            row[j+1].text = clean_text(cell_text)

    out_path = results_folder.with_suffix(".docx")
    doc.save(out_path)
    print(f"✅ Written: {out_path.resolve()}")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python write_results_table.py /path/to/Results_folder/")
        sys.exit(1)
    write_results_table(sys.argv[1])
